
from openpyxl import load_workbook

from docx import Document
from docx.shared import Inches

data = []

def main():
    workbook = load_workbook(filename="data.xlsx")
    sheet = workbook.active
    print(sheet.cell(row=3,column=2).value)
    for col in sheet.iter_cols(min_row = 1, max_row = 30, min_col = 1, max_col = 30,values_only=True):
        data.append(col)
        #print(col)
    counter =0
    for d in data:
            
        if counter == 0:
            counter += 1
            continue
        print(d)
        print(f"Filename {d[0]}")
        write_file(d[0],d)
        counter += 1
        
def write_file(filename,data):
    document = Document()
    document.add_heading(filename,0)

    for i in range(1,30):
        if str(data[i]).strip() != "None":
            print("here", data[i])
            document.add_paragraph(data[i],style = "List Bullet")
    print("file saved")
    document.save(f"out/{filename.strip()}.pdf")
        
if __name__ == "__main__":
    main()

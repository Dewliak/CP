
from openpyxl import load_workbook

from docx import Document
from docx.shared import Inches

import random 

data = []

def main():
    document = Document()


    workbook = load_workbook(filename="data.xlsx")
    sheet = workbook.active
    print(sheet.cell(row=3,column=2).value)
    for col in sheet.iter_cols(min_row = 1, max_row = 30, min_col = 1, max_col = 30,values_only=True):
        data.append(col)

    counter =0

    for d in data:
            
        if counter == 0:
            counter += 1
            continue

        random.shuffle(d)

        print(d)
        print(f"Filename {d[0]}")

        document.add_heading(filename,0)

        for i in range(1,30):

            if str(d[i]).strip() != "None":
                print("here", d[i])
                document.add_paragraph(d[i],style = "List Bullet")

        document.add_page_break()

        counter += 1

    document.save(f"out/osztaly.pdf")
        

        
if __name__ == "__main__":
    main()
